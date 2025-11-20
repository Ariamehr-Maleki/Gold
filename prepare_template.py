import argparse
import os
import logging
from docx import Document

# Setup basic logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

def prepare_template(original_path, output_path):
    """
    Reads the 'human' DOCX template (from the text you provided) and 
    replaces generic placeholders (e.g. [value], [growth rate]) with 
    specific, unique keys (e.g. [export_val], [tm_cagr]) that the 
    Orchestrator uses.
    """
    if not os.path.exists(original_path):
        logging.error(f"FATAL: Original template not found at '{original_path}'")
        return

    # --- MAPPING STRATEGY ---
    # The order is critical. We must match longer, specific sentences first
    # to handle context collision (e.g., preventing [value] for 'Imports' being 
    # replaced by [value] for 'Unit Price').

    replacements = [
        # --- Header & General ---
        ("[Product]", "[product_name]"),
        ("[product]", "[product_name]"), 
        ("[Target Market]", "[target_market]"),
        ("[Target market]", "[target_market]"),
        ("[Target market]’s", "[target_market]’s"), # Handle smart quote possession
        ("[target market]", "[target_market]"),
        ("[Your Country]", "[your_country]"),
        ("[Your country]", "[your_country]"),
        ("[your country]", "[your_country]"),
        ("[Month Year]", "[month_year]"),
        ("[00.00.00]", "[hs_code]"),

        # --- Introduction Section ---
        # Rank
        ("Rank in World for Imports of this Product\t[rank]", "Rank in World for Imports of this Product\t[rank]"), 
        # Total Exports (The first [value] collision)
        ("Total exports from [your_country] in [year] to the world\tUSD [value]", "Total exports from [your_country] in [year_latest] to the world\tUSD [export_val]"),

        # --- Opportunity Summary (Handling repeated [latest data]) ---
        ("Capital City\t[enter city] *", "Capital City\t[city_name] *"),
        ("Population\t[latest data] *", "Population\t[population] *"),
        ("GDP per capita\t[latest data] *", "GDP per capita\t[gdp] *"),
        ("Currency\t[latest data] *", "Currency\t[currency] *"),
        ("Languages\t[latest data] *", "Languages\t[languages] *"),

        # --- Size of the Market ---
        # Collision: "imported USD [value]" vs Unit Value
        ("imported USD [value] of [product_name]s from the world", "imported USD [tm_global_import_val] of [product_name] from the world"),
        ("represented [share] %", "represented [tm_world_share]%"),
        ("imported USD [value] of the product from [your_country]", "imported USD [yc_market_import_val] of the product from [your_country]"),
        ("means [your_country] has a [share] %", "means [your_country] has a [yc_share_pct]%"),
        
        # --- Growth of the Market ---
        # Collision: "grew by [growth rate]" appears 3 times with different meanings
        # 1. Market Growth
        ("from the world of [product_name] grew by [growth rate] %", "from the world of [product_name] grew by [tm_cagr]%"),
        ("performance was [better than / worse]", "performance was [compare_growth]"),
        # 2. World Growth
        ("growth in imports of [product_name], which grew by [growth rate] %", "growth in imports of [product_name], which grew by [world_cagr]%"),
        ("imports of [product_name] has been [increasing / decreasing]", "imports of [product_name] has been [share_trend]"),
        ("growth rate was [sustained / not sustained]", "growth rate was [sustained_trend]"),
        # 3. Last Year Growth
        ("market [growing / contracting] by [growth rate] %", "market [last_year_trend] by [last_year_growth]%"),
        # 4. Your Country Growth
        ("[target_market]’s imports from [your_country] grew by [growth rate] %", "[target_market]’s imports from [your_country] grew by [yc_cagr]%"),
        ("[your_country] [gained / lost] market share", "[your_country] [share_change_trend] market share"),

        # --- Unit Value (High collision zone for [value]) ---
        # 1. Target Market Average
        ("imports of [product_name] in [year] was [value] USD/ [unit]", "imports of [product_name] in [year_latest] was [uv_tm] USD/[unit]"),
        # 2. World Comparison
        ("more than/less than]", "uv_compare]"), # Partial matching for safety
        ("world unit value for the product of [value] USD/ [unit]", "world unit value for the product of [uv_world] USD/[unit]"),
        ("unit value has been [appreciating/depreciating]", "unit value has been [uv_trend]"),
        # 3. Your Country
        ("paid by [target_market] to [your_country] was [value] USD/ [unit]", "paid by [target_market] to [your_country] was [uv_yc] USD/[unit]"),
        ("[appreciated/depreciated]", "[uv_trend_yc]"),
        # 4. High/Low Range
        ("highest unit value of [value] USD/ [unit] paid to [country X]", "highest unit value of [uv_high] USD/[unit] paid to [supplier_high]"),
        ("lowest unit value of [value] USD/ [unit] paid to [country Y]", "lowest unit value of [uv_low] USD/[unit] paid to [supplier_low]"),
        ("[rather heterogeneous/somewhat homogeneous]", "[heterogeneity]"),

        # --- Competition ---
        ("market for [product_name] is [not concentrated / moderately concentrated / concentrated]", "market for [product_name] is [concentration_level]"),
        ("exporters, [supplier 1]; [supplier 2]; [supplier 3]", "exporters, [s1_name]; [s2_name]; [s3_name]"),
        # Pie Chart Captions (complex matching might fail, safer to rely on python text insertion)
        ("having market shares of [value] %, [value] % and [value] %", "having market shares of [s1_share]%, [s2_share]% and [s3_share]%"),
        ("include [supplier A]; [supplier B]; [supplier C]", "include [suppliers_gaining_share]"),
        ("from [your region] include: [supplier X]; [supplier Y]; [supplier Z]", "from [your region] include: [regional_competitors]"),

        # --- Market Access & Tables ---
        ("Table row code: [00.00.00.00.00]", "N/A - This is usually handled by Orchestrator insertion logic"),
        
        # Placeholders specifically for images (Orchestrator looks for these phrases)
        ("[Good quality picture \nof the product]", "[img_product_placeholder]"),
        ("[World map or regional map highlighting \ntarget market’s country location]", "[img_map_placeholder]"),
        ("[Line graph of target market’s total imports of the selected product in the last 5-10 years]", "[img_line_graph_placeholder]"),
        ("[Pie graph showing last year’s market shares of main suppliers of the selected product to the target market]", "[img_pie_chart_placeholder]")
    ]

    doc = Document(original_path)

    # Perform Replacement on Paragraphs
    for p in doc.paragraphs:
        # Iterate through our prioritized list
        for old_txt, new_txt in replacements:
            if old_txt in p.text:
                # direct string replacement works best for standard docx text
                p.text = p.text.replace(old_txt, new_txt)

    # Perform Replacement on Tables (Recursive)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for old_txt, new_txt in replacements:
                        if old_txt in p.text:
                            p.text = p.text.replace(old_txt, new_txt)

    # Post-Processing: Ensure the Instruction page content is irrelevant 
    # (The orchestrator deals with deletion, but we make sure formatting is clean here)

    logging.info(f"Template updates applied.")
    doc.save(output_path)
    logging.info(f"✅ Success! Machine-ready template saved to: '{output_path}'")

if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument("--original", required=True, help="Path to the unmodified docx containing [value], [product] etc.")
    parser.add_argument("--output", required=True, help="Path to save the prepared template.")
    args = parser.parse_args()

    prepare_template(args.original, args.output)