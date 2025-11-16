import argparse
import json
import os
import re
from datetime import datetime
from docx import Document
from docx2pdf import convert
import logging

# Setup basic logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')

class FactsheetGenerator:
    """
    Generates a PDF factsheet by populating a DOCX template with data from a JSON file.
    """
    def __init__(self, json_path, template_path, output_path):
        self.json_path = json_path
        self.template_path = template_path
        self.output_path = output_path
        self.data = None
        self.doc = None
        self.context = {}

    def _load_data(self):
        """Loads and validates the source JSON file."""
        try:
            with open(self.json_path, 'r', encoding='utf-8') as f:
                self.data = json.load(f)
            logging.info(f"Successfully loaded data from '{self.json_path}'")
        except FileNotFoundError:
            logging.error(f"FATAL: JSON file not found at '{self.json_path}'")
            raise
        except json.JSONDecodeError:
            logging.error(f"FATAL: Could not decode JSON from '{self.json_path}'. File may be corrupt.")
            raise

    def _create_context(self):
        """
        Creates a flat key-value dictionary (the 'context') for easy replacement.
        This is where all calculations and data transformations happen.
        """
        def get(path, default="N/A"):
            """Helper to safely get a value from the nested JSON data."""
            keys = path.split('.')
            val = self.data
            for key in keys:
                try:
                    val = val[key] if isinstance(val, dict) else val[int(key)]
                except (KeyError, TypeError, IndexError, ValueError):
                    return default
            return val if val not in [None, ''] else default

        def format_num(value):
            """Helper to format numbers with commas."""
            try: return f"{int(value):,}"
            except (ValueError, TypeError): return str(value)

        # --- Build Context Dictionary ---
        ctx = {}
        # Basic Info
        ctx['product_name'] = get('market_access_conditions.customs_tariffs.0.Product description', 'the selected product')
        ctx['target_market'] = get('meta.importing_country')
        ctx['your_country'] = get('meta.exporting_country')
        ctx['month_year'] = datetime.now().strftime("%B %Y")
        ctx['hs_code'] = get('meta.product_hs6')
        ctx['year'] = get('market_concentration.top_5_suppliers.0.year', datetime.now().year) # Fallback to current year

        # Opportunity Summary
        ctx['your_total_exports'] = "Data not available"  # This data is not in the current JSON
        ctx['world_rank'] = get('market_size_and_potential.target_market_world_rank')
        
        # --- Size of Market ---
        tm_imports = get('market_size_and_potential.total_market_value_usd', 0)
        world_imports = get('importer_ranking.world_total_imports_usd', 0)
        # Your country's value imported BY the target market
        imports_from_yc = next((s['value_usd'] for s in get('market_concentration.top_5_suppliers', []) if s.get('name') == ctx['your_country']), 0)

        ctx['tm_total_imports'] = format_num(tm_imports)
        ctx['tm_share_of_world'] = round((tm_imports / world_imports) * 100, 2) if world_imports else 0
        ctx['imports_from_yc'] = format_num(imports_from_yc)
        ctx['yc_share_of_tm'] = get('competitive_landscape.your_country_market_share_pct', 0)

        # --- Growth of Market ---
        tm_cagr = get('market_size_and_potential.market_growth_cagr_5y_pct', 0)
        tm_last_growth = get('market_analysis.market_growth_last_year_pct', 0) # Assumes this key exists, if not use N/A
        yc_growth = next((s['growth_cagr_pct'] for s in get('market_concentration.top_5_suppliers', []) if s.get('name') == ctx['your_country']), 0)
        
        ctx['tm_growth_cagr'] = tm_cagr
        ctx['world_growth_cagr'] = tm_cagr 
        ctx['tm_vs_world_growth'] = "better" if tm_cagr > 0 else "worse"
        ctx['tm_share_trend'] = "increasing" if tm_cagr > 0 else "decreasing"
        ctx['tm_last_year_trend'] = "growing" if tm_last_growth > 0 else "contracting"
        ctx['tm_last_year_growth'] = tm_last_growth
        ctx['yc_growth_cagr'] = yc_growth
        ctx['yc_share_change'] = "gained" if get('competitive_landscape.your_country_market_share_pct', 0) > 0 else "lost"

        # --- Competition ---
        ctx['concentration'] = get('market_concentration.concentration_level')
        top_suppliers = get('market_concentration.top_5_suppliers', [])
        for i in range(3):
            s = top_suppliers[i] if i < len(top_suppliers) else {}
            ctx[f's{i+1}_name'] = s.get('name', 'N/A')
            ctx[f's{i+1}_share'] = s.get('market_share_pct', 'N/A')

        ctx['suppliers_gaining_share'] = ", ".join(get('market_concentration.suppliers_gaining_share', ['N/A']))
        ctx['regional_competitors'] = ", ".join([s['name'] for s in get('competitive_landscape.regional_competitors', []) if s.get('name')] or ["N/A"])

        # These placeholders were missing from the script but are in your doc
        ctx['value'] = "N/A"
        ctx['unit'] = "N/A"
        ctx['period'] = "N/A"
        ctx['your_region'] = "Asia" # Hardcoded as an example, this isn't in the JSON

        self.context = ctx

    def _replace_placeholders(self):
        """Finds and replaces all [key] style placeholders throughout the document."""
        placeholder_regex = re.compile(r'\[([^\[\]]+)\]')

        def substitute_in_paragraph(p):
            """Replaces all placeholders in a single paragraph's text."""
            # We need to do this in a loop because a replacement might contain another placeholder
            for _ in range(10): # Max 10 iterations to prevent infinite loops
                found_placeholders = re.findall(placeholder_regex, p.text)
                if not found_placeholders:
                    break
                for key in found_placeholders:
                    # Look for the key in context, if not found, keep the original placeholder
                    value = str(self.context.get(key.strip(), f"[{key}]"))
                    p.text = p.text.replace(f"[{key}]", value, 1) # Replace one at a time

        def process_element(element):
            for p in element.paragraphs:
                substitute_in_paragraph(p)
            for table in element.tables:
                for row in table.rows:
                    for cell in row.cells:
                        process_element(cell) # Recurse into cells
        
        logging.info("Replacing placeholders in document...")
        process_element(self.doc)

    def _populate_dynamic_tables(self):
        """Populates tables that require dynamic row creation."""
        companies = self.data.get('competitive_landscape', {}).get('potential_importers', [])
        if not companies: return

        for table in self.doc.tables:
            if len(table.rows) > 1 and '[company_loop_start]' in table.rows[1].cells[0].text:
                placeholder_row = table.rows[1]
                for company in companies[:5]: # Add top 5
                    row_cells = table.add_row().cells
                    row_cells[0].text = company.get('name', '')
                    row_cells[1].text = company.get('city', '')
                    row_cells[2].text = company.get('website', '')
                placeholder_row._element.getparent().remove(placeholder_row._element)
                logging.info(f"Populated companies table with {len(companies[:5])} entries.")
                break

    def generate(self):
        """Runs the complete factsheet generation process."""
        try:
            self._load_data()
            self.doc = Document(self.template_path)
            self._create_context()
            self._replace_placeholders()
            self._populate_dynamic_tables()

            temp_docx_path = os.path.join(os.path.dirname(self.output_path), "temp_filled_factsheet.docx")
            self.doc.save(temp_docx_path)
            logging.info(f"Saved filled Word document to: {temp_docx_path}")

            logging.info("Converting to PDF...")
            convert(temp_docx_path, self.output_path)
            logging.info(f"✅ Success! Final PDF factsheet saved to: {self.output_path}")

        except Exception as e:
            logging.critical(f"An unexpected error occurred: {e}", exc_info=True)
        finally:
            # Clean up temp file
            if 'temp_docx_path' in locals() and os.path.exists(temp_docx_path):
                os.remove(temp_docx_path)

if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Generate a PDF factsheet from a JSON file and a DOCX template.")
    parser.add_argument("--json-file", required=True, help="Path to the final_report.json file.")
    parser.add_argument("--template", required=True, help="Path to the prepared .docx template.")
    parser.add_argument("--output", required=True, help="Path to save the final .pdf report.")
    args = parser.parse_args()
    
    os.makedirs(os.path.dirname(args.output), exist_ok=True)
    
    generator = FactsheetGenerator(args.json_file, args.template, args.output)
    generator.generate()