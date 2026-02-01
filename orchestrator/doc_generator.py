# orchestrator/doc_generator.py

import json
import os
import logging
import re
from docx import Document
from docx.shared import Inches

# Configuration
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

class FactsheetGenerator:
    def __init__(self, json_path, output_dir):
        # We now expect json_path to point to the Hierarchical 'factsheet_data.json'
        self.json_path = json_path
        self.output_dir = output_dir
        self.data = self._load_data()
        self.replacements = {}

    def _load_data(self):
        if not os.path.exists(self.json_path):
            raise FileNotFoundError(f"JSON file not found at {self.json_path}")
        with open(self.json_path, 'r', encoding='utf-8') as f:
            return json.load(f)

    def _get(self, path, default="N/A"):
        """Helper to navigate the Hierarchical JSON safely."""
        keys = path.split('.')
        val = self.data
        for key in keys:
            if isinstance(val, dict):
                val = val.get(key)
            elif isinstance(val, list):
                try:
                    val = val[int(key)]
                except (IndexError, ValueError):
                    return default
            else:
                return default
            
            if val is None:
                return default
        return str(val)

    def build_mappings(self):
        """
        Maps the Hierarchical JSON keys to the specific Placeholders used in the Word Template.
        """
        # Shortcuts for cleaner code
        h = "Header"
        cov = "Cover"
        intro = "Introduction"
        size = "Size_of_the_Market"
        growth = "Growth_of_the_Market"
        uv = "Unit_Value"
        comp = "Competition"
        ma = "Market_Access"

        self.replacements = {
            # --- Header & Intro ---
            "[Name of Country]": self._get(f"{h}.Target_Market"),
            "[Target Market]": self._get(f"{h}.Target_Market"),
            "[Target market]": self._get(f"{h}.Target_Market"),
            "[Your country]": self._get(f"{intro}.Your_Country"),
            "[Your Country]": self._get(f"{intro}.Your_Country"),
            "[product]": self._get(f"{intro}.Product_Name"),
            "HS 281511": self._get(f"{intro}.HS_Code"), # Fallback for hardcoded template text
            
            # --- Summary (Derived from Trade Overview) ---
            "[world_rank]": self._get(f"Trade_Overview.Rank_in_World_For_Imports_Of_This_Product"),
            "[capital_city]": self._get(f"Opportunity_Summary.Capital_City"),
            "[population]": self._get(f"Opportunity_Summary.Population"),
            "[currency]": self._get(f"Opportunity_Summary.Currency"),
            
            # --- Size ---
            "[tm_total_imports]": self._get(f"{size}.Target_Market_Imported_Value_From_World_USD"),
            "[tm_share_of_world]": self._get(f"{size}.World_Import_Share_Percent"),
            "[imports_from_yc]": self._get(f"{size}.Target_Market_Imported_Value_From_Your_Country_USD"),
            "[yc_share_of_tm]": self._get(f"{size}.Your_Country_Share_Of_Target_Imports_Percent"),
            
            # --- Growth ---
            "[tm_growth_cagr]": self._get(f"{growth}.Five_Year_Growth_Rate_Target_Market_Percent"),
            "[tm_vs_world_growth]": self._get(f"{growth}.Performance_Compared_To_World"),
            "[world_growth_cagr]": self._get(f"{growth}.World_Imports_Growth_Rate_Percent"),
            "[tm_share_trend]": self._get(f"{growth}.Target_Market_Share_Trend"),
            "[sustained / not sustained]": self._get(f"{growth}.Recent_Growth_Sustained_or_Not"),
            "[tm_last_year_trend]": self._get(f"{growth}.Recent_Growth_Direction"),
            "[tm_last_year_growth]": self._get(f"{growth}.Recent_Growth_Rate_Percent"),
            "[yc_growth_cagr]": self._get(f"{growth}.Five_Year_Growth_Rate_Your_Country_Percent"),
            "[yc_share_change]": self._get(f"{growth}.Your_Country_Market_Share_Change"),
            
            # --- Unit Values ---
            "58,630 USD/ unit": f"{self._get(f'{uv}.Target_Market_Avg_Unit_Value.Value_USD')} {self._get(f'{uv}.Target_Market_Avg_Unit_Value.Unit')}",
            "[more than/less than]": self._get(f"{uv}.Comparison_To_World_Unit_Value_Statement"),
            "[appreciated/depreciated]": self._get(f"{uv}.Target_Market_Unit_Value_Trend"),
            "[appreciating/depreciating].": self._get(f"{uv}.World_Unit_Value_Trend") + ".", 
            "[country X]": self._get(f"{uv}.Top_Ten_Suppliers_Unit_Value_Range.Highest_Unit_Value.Country"),
            "[country Y]": self._get(f"{uv}.Top_Ten_Suppliers_Unit_Value_Range.Lowest_Unit_Value.Country"),
            "[rather heterogeneous/somewhat homogeneous]": self._get(f"{uv}.Top_Ten_Suppliers_Unit_Value_Range.Market_Heterogeneity_Statement"),
            
            # --- Competition ---
            "[not concentrated / moderately concentrated / concentrated]": self._get(f"{comp}.Market_Concentration_Level"),
            "[your_region]": "your region", # Static placeholder
        }

    def _process_placeholders(self, doc):
        """Iterates through paragraphs and tables to replace text."""
        def replace_in_text(text):
            for key, val in self.replacements.items():
                if key in text:
                    # Robust replacement handling strings
                    text = text.replace(key, str(val))
            return text

        for p in doc.paragraphs:
            if "[" in p.text:
                p.text = replace_in_text(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if "[" in p.text:
                            p.text = replace_in_text(p.text)

    def _fill_tariff_table(self, doc):
        """Fills the Tariff table from the Market Access section."""
        target_table = None
        for t in doc.tables:
            if not t.rows: continue
            # Identify table by checking header for "tariff" or "code"
            header_text = t.rows[0].cells[0].text.lower()
            if "tariff" in header_text or "code" in header_text:
                target_table = t
                break
        
        if not target_table:
            return

        tariffs = self.data.get("Market_Access", {}).get("Tariff_Table", [])
        
        # We fill up to 5 rows, assuming row 0 is header
        for i in range(5):
            # Add row if needed
            if i + 1 >= len(target_table.rows):
                if i < len(tariffs): target_table.add_row()
                else: break 
            
            row = target_table.rows[i + 1]
            if i < len(tariffs):
                item = tariffs[i]
                row.cells[0].text = str(item.get("National_Tariff_Line_Code", ""))
                if len(row.cells) > 1: row.cells[1].text = str(item.get("Product_Description", ""))[:50]
                if len(row.cells) > 2: row.cells[2].text = str(item.get("MFN_or_General_Tariff", ""))
                # Add other columns as needed matching template columns
            else:
                # Clear unused rows
                for cell in row.cells: cell.text = ""

    def process_document(self, template_path):
        logger.info(f"Loading template from {template_path}...")
        doc = Document(template_path)
        
        # 1. Build the mapping dictionary
        self.build_mappings()
        
        # 2. Replace simple placeholders
        self._process_placeholders(doc)
        
        # 3. Fill specific tables
        self._fill_tariff_table(doc)
        # Note: You can add _fill_importers_table here similar to tariff table if needed
        
        # 4. Save
        output_filename = f"Factsheet_{self._get('Header.Product')}_{self._get('Header.Target_Market')}.docx"
        output_filename = re.sub(r'[\\/*?:"<>|]', "", output_filename) # Sanitize filename
        save_path = os.path.join(self.output_dir, output_filename)
        
        doc.save(save_path)
        logger.info(f"Report generated successfully: {save_path}")

def run(json_path, template_path, output_dir):
    # Important: We assume json_path passed here is the NEW factsheet_data.json
    # If the orchestrator passes final_report.json, we swap it here.
    
    if "final_report.json" in json_path:
        # Switch to the structured factsheet data file
        factsheet_json_path = os.path.join(os.path.dirname(json_path), 'factsheet_data.json')
        if os.path.exists(factsheet_json_path):
            json_path = factsheet_json_path
            logger.info(f"Switched input to structured data: {json_path}")

    gen = FactsheetGenerator(json_path, output_dir)
    gen.process_document(template_path)

if __name__ == "__main__":
    pass